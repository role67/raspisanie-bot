import pandas as pd
import requests
import re
from io import BytesIO
from docx import Document
import random
import os
from pathlib import Path
import threading
import logging

SCHEDULE_URL = "https://www.nkptiu.ru/doc/raspisanie/raspisanie.xls"
REPLACEMENTS_URL = "https://www.nkptiu.ru/doc/raspisanie/zameni.docx"

def load_user_agents():
    """Загружает User-Agent'ы из файлов"""
    agents = {
        'windows': [],
        'mac': [],
        'ios': [],
        'android': []
    }
    
    base_path = Path(__file__).parent.parent / 'useragents'
    
    # Загружаем по 100 агентов каждого типа
    for platform in agents.keys():
        file_path = base_path / f"{platform}.txt"
        if file_path.exists():
            with open(file_path, 'r', encoding='utf-8') as f:
                # Берем первые 100 строк, пропуская пустые
                agents[platform] = [line.strip() for line in f if line.strip()][:100]
    
    return agents


# Загружаем User-Agent'ы при импорте модуля
USER_AGENTS = load_user_agents()

# Глобальный кэш расписания и блокировка
_schedule_cache = None
_schedule_cache_lock = threading.Lock()
_schedule_cache_hash = None

# --- Новый парсер строки расписания ---
def split_subject_teacher(cell: str):

    cell = cell.strip()
    pattern = re.compile(
        r"""
        ^(?P<subject>[А-Яа-яA-Za-zЁё .\-]+?)\s*
        (?:\((?P<subgroup>\dп)\))?\s*
        (?P<teacher>[А-ЯЁ][а-яё]+\s[А-ЯЁ]\.[А-ЯЁ]\.)?\s*
        (?P<room>\d{2,4})?
        $
        """, re.VERBOSE)
    match = pattern.match(cell)
    if match:
        subject = (match.group('subject') or '').strip()
        subgroup = (match.group('subgroup') or '').strip()
        teacher = (match.group('teacher') or '').strip()
        room = (match.group('room') or '').strip()
        return subject, teacher, room, subgroup
    # fallback: попытка вытащить кабинет
    room_match = re.search(r"(\d{2,4})$", cell)
    room = room_match.group(1) if room_match else ''
    teacher_match = re.search(r"([А-ЯЁ][а-яё]+\s[А-ЯЁ]\.[А-ЯЁ]\.)", cell)
    teacher = teacher_match.group(1) if teacher_match else ''
    subgroup_match = re.search(r"\((\dп)\)", cell)
    subgroup = subgroup_match.group(1) if subgroup_match else ''
    subject = cell
    for part in [teacher, room, f"({subgroup})"]:
        if part:
            subject = subject.replace(part, '').strip()
    subject = re.sub(r"\s+", " ", subject)
    return subject, teacher, room, subgroup

# --- Форматирование расписания дня ---
def format_day_schedule(group_lessons, day, replacements=None, date_str=None, last_update=None):
    """
    Форматирует расписание дня для вывода пользователю
    """
    try:
        if not isinstance(group_lessons, dict):
            return "❌ Ошибка в формате данных расписания"
        if not day:
            return "❌ Не указан день недели"
        if replacements is not None and not isinstance(replacements, list):
            replacements = None
        if date_str and not isinstance(date_str, str):
            try:
                date_str = str(date_str)
            except:
                date_str = None
        from datetime import datetime
        try:
            from zoneinfo import ZoneInfo
            tz_msk = ZoneInfo("Europe/Moscow")
        except ImportError:
            from pytz import timezone
            tz_msk = timezone("Europe/Moscow")
        now_msk = datetime.now(tz_msk)
        weekday = now_msk.weekday()
        iso_week = now_msk.isocalendar().week
        if weekday == 6 and now_msk.hour >= 0:
            week_number = 1 if (iso_week + 1) % 2 != 0 else 2
        else:
            week_number = 2 if iso_week % 2 == 0 else 1
    except Exception:
        return "❌ Ошибка при обработке данных расписания"
    from .lesson_times import LESSON_TIMES, WEEKDAY_TIMES, SATURDAY_TIMES
    day_map = {
        'Понедельник': 'Понедельник',
        'Вторник': 'Вторник',
        'Среда': 'Среда',
        'Четверг': 'Четверг',
        'Пятница': 'Пятница',
        'Суббота': 'Суббота'
    }
    if day == 'Понедельник':
        times_dict = LESSON_TIMES
    elif day == 'Суббота':
        times_dict = SATURDAY_TIMES
    else:
        times_dict = WEEKDAY_TIMES
    try:
        if date_str:
            lines = [f"📅 {date_str} | {day_map.get(day, day)}\n"]
        else:
            lines = [f"📅 {day_map.get(day, day)}\n"]
        if not group_lessons or day not in group_lessons:
            lines.append("\n❌ Расписание на этот день не найдено")
            return '\n'.join(lines)
        lessons = group_lessons.get(day, {}).get(week_number, [])
        if not isinstance(lessons, list):
            lines.append("\n❌ Ошибка в формате расписания")
            return '\n'.join(lines)
        for idx, lesson in enumerate(lessons, 1):
            try:
                if not isinstance(lesson, dict):
                    continue
                subject = lesson.get('subject', '').strip()
                teacher = lesson.get('teacher', '').strip()
                room = lesson.get('room', '').strip()
                time = lesson.get('time', '').strip()
                if not subject or subject == "-----":
                    continue
            except Exception:
                continue
            time_str = times_dict.get(time, time)
            lesson_str = f"{idx}️⃣ {subject} | {time_str}"
            lines.append(lesson_str)
            if teacher:
                lines.append(f"👤 {teacher}")
            if room:
                lines.append(f"🚪 {room}")
            lines.append("")
    except Exception:
        pass
    try:
        if replacements:
            if not isinstance(replacements, list):
                lines.append("❌ Ошибка при обработке замен")
            else:
                lines.append("🔄 Замены")
                for idx, rep in enumerate(replacements, 1):
                    try:
                        if not isinstance(rep, dict):
                            continue
                        rep_subject = rep.get('subject', '').strip()
                        rep_lesson = rep.get('lesson', '').strip()
                        rep_room = rep.get('room', '').strip()
                        rep_teacher = rep.get('teacher', '').strip()
                        if not rep_subject or not rep_lesson:
                            continue
                        lines.append(f"📚 {rep_subject} вместо {rep_lesson}")
                        if rep_teacher:
                            lines.append(f"👤 {rep_teacher}")
                        if rep_room:
                            lines.append(f"🚪 Каб. {rep_room}")
                        lines.append("")
                    except Exception:
                        continue
    except Exception:
        lines.append("❌ Ошибка при обработке замен")
    try:
        if last_update:
            try:
                update_time = last_update.strftime('%d.%m.%Y %H:%M')
                lines.append(f"🕒 Обновлено: {update_time}")
            except Exception:
                lines.append("🕒 Время обновления недоступно")
        lines.append(f"📅 {week_number} неделя")
    except Exception:
        pass
    try:
        return '\n'.join(lines)
    except Exception:
        return "❌ Ошибка при формировании расписания"

def get_random_headers():
    """Возвращает случайный User-Agent и базовые заголовки"""
    # Выбираем платформу с разными весами
    platform = random.choices(
        ['windows', 'mac', 'ios', 'android'],
        weights=[0.4, 0.3, 0.2, 0.1]  # 40% Windows, 30% Mac, 20% iOS, 10% Android
    )[0]
    
    # Получаем список агентов для выбранной платформы
    agents = USER_AGENTS.get(platform, [])
    
    # Если список пуст, пробуем взять случайный агент из любой доступной платформы
    if not agents:
        all_agents = []
        for platform_agents in USER_AGENTS.values():
            all_agents.extend(platform_agents)
        user_agent = random.choice(all_agents) if all_agents else (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/118.0.0.0 Safari/537.36"
        )
    else:
        user_agent = random.choice(agents)
    
    return {
        "User-Agent": user_agent,
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
        "Accept-Language": "ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7",
        "Accept-Encoding": "gzip, deflate, br",
        "Connection": "keep-alive",
        "Cache-Control": "no-cache",
        "Pragma": "no-cache"
    }


def process_subject_and_teacher(value):
    """Разделяет предмет и преподавателя из одной строки"""
    try:
        if not value or pd.isna(value) or str(value).strip().lower() == 'nan':
            return '', ''
    except Exception as e:
        print(f"Ошибка при проверке значения: {e}")
        return '', ''
    
    value = str(value).strip()
    parts = value.split(')')  # Разделяем по скобке, если есть
    
    # Если есть скобки (например: "МДК.04.01 (Болдовская)")
    if len(parts) > 1:
        subject = parts[0].strip() + ')'
        teacher = parts[1].strip()
        return subject, teacher
    
    # Если есть явное указание языка
    if '(нем)' in value.lower() or '(англ)' in value.lower():
        parts = value.split()
        subject_parts = []
        teacher_parts = []
        found_lang = False
        
        for part in parts:
            if '(нем)' in part.lower() or '(англ)' in part.lower():
                subject_parts.append(part)
                found_lang = True
            elif found_lang:
                teacher_parts.append(part)
            else:
                subject_parts.append(part)
                
        return ' '.join(subject_parts), ' '.join(teacher_parts)
    
    return value, ''

def process_teacher_and_room(value):
    """Разделяет учителя и кабинет из строки"""
    try:
        if not value or pd.isna(value) or str(value).strip().lower() == 'nan':
            return '', ''
    except Exception as e:
        print(f"Ошибка при проверке значения: {e}")
        return '', ''
        
    value = str(value).strip()
    
    # Специальная обработка общежития
    if value.lower() in ['общ', 'общ.', 'общага']:
        return '', 'Общежитие'
        
    # Если есть номер через дефис (например 201-4), это кабинет
    if '-' in value and any(c.isdigit() for c in value) and len(value) <= 7:
        return '', f"Каб. {value}"
        
    # Если это просто номер, это тоже кабинет
    if value.isdigit() and len(value) <= 4:
        return '', f"Каб. {value}"
    
    # Признаки преподавателя:
    # 1. Фамилия с инициалами (Иванов И.И.)
    if ' ' in value and any(c == '.' for c in value) and value.split()[0].istitle():
        return value, ''
    
    # 2. Просто фамилия с большой буквы
    if value.istitle() and value.isalpha() and len(value) > 3:
        return value, ''
        
    # 3. Фамилия с пробелами без точек (Иванова Мария)
    if ' ' in value and all(word.istitle() for word in value.split()) and value.replace(' ', '').isalpha():
        return value, ''
    
    # Иначе считаем что это не преподаватель и не кабинет
    return '', ''

def fetch_schedule():
    """Получает и парсит основное расписание"""
    global _schedule_cache, _schedule_cache_lock, _schedule_cache_hash
    try:
        with _schedule_cache_lock:
            headers = get_random_headers()
            try:
                resp = requests.get(SCHEDULE_URL, headers=headers, timeout=30)
                resp.raise_for_status()
                if resp.status_code != 200 or len(resp.content) < 1000:
                    logging.error(f"Ошибка при получении файла расписания: статус={resp.status_code}, длина={len(resp.content)}")
                    return {}
                file_hash = hash(resp.content)
                if _schedule_cache is not None and _schedule_cache_hash == file_hash:
                    logging.info("Возвращаем кэш расписания")
                    return _schedule_cache.copy() if isinstance(_schedule_cache, dict) else {}
                xls = BytesIO(resp.content)
            except Exception as e:
                logging.error(f"Ошибка при загрузке расписания: {e}")
                return {}
        try:
            try:
                df = pd.read_excel(xls, engine='xlrd', na_values=[''])
            except Exception as e1:
                xls.seek(0)
                try:
                    df = pd.read_excel(xls, engine='openpyxl', na_values=[''])
                except Exception as e2:
                    logging.error(f"Ошибка чтения xls: xlrd={e1}, openpyxl={e2}")
                    return {}
            if df.empty or len(df.columns) < 3:
                logging.error(f"DataFrame пустой или мало колонок: shape={df.shape}, columns={df.columns}")
                return {}
            logging.info(f"DataFrame загружен: shape={df.shape}, columns={list(df.columns)}")
        except Exception as e:
            logging.error(f"Ошибка при обработке DataFrame: {e}")
            return {}
        practice_rows = df[df.iloc[:, 0] == "ПРАКТИКИ"].index
        practice_data = {}
        
        if len(practice_rows) > 0:
            practice_start = practice_rows[0]
            # Читаем практики после строки "ПРАКТИКИ"
            for idx, row in df.iloc[practice_start+1:].iterrows():
                try:
                    if pd.notna(row[0]) and str(row[0]).strip():
                        group = str(row[0]).strip()
                        # Получаем все непустые значения из строки
                        practice_values = []
                        for i in range(1, len(row)):
                            if pd.notna(row[i]) and str(row[i]).strip():
                                practice_values.append(str(row[i]).strip())
                        
                        # Если нашли хотя бы одно значение и это не заголовок "ПРАКТИКИ"
                        if practice_values and group != "ПРАКТИКИ":
                            practice_info = " ".join(practice_values)
                            if group and practice_info:
                                practice_data[group] = practice_info
                                logging.debug(f"Добавлена информация о практике для группы {group}")
                except (IndexError, TypeError, AttributeError) as e:
                    logging.error(f"Ошибка при обработке строки практики {idx}: {e}")
        
    # Логирование убрано для оптимизации
        
        # Заполняем пропуски времени (интервала)
        if 'Интервал' in df.columns:
            df['Интервал'] = df['Интервал'].ffill()
            logging.debug("Заполнены пропуски в столбце 'Интервал'")

        # Импортируем времена пар
        from .lesson_times import LESSON_TIMES, WEEKDAY_TIMES, SATURDAY_TIMES

        schedule_data = {}
        
        # Определяем колонки групп по шаблону: буквы+дефис+цифры (например "ИСП-21")
        group_cols = [col for col in df.columns 
                     if isinstance(col, str) and 
                     '-' in col and 
                     any(c.isalpha() for c in col) and 
                     any(c.isdigit() for c in col)]
        
        # if not group_cols:
        #     return {}

        for group_col in group_cols:
            schedule_data[group_col] = {}
            try:
                if group_col in practice_data:
                    practice_info = practice_data[group_col]
                    # Проверяем что информация о практике не пустая
                    if practice_info and practice_info.strip():
                        schedule_data[group_col] = {
                            'practice': [{
                                'is_practice': True,
                                'practice_info': practice_info.strip()
                            }],
                            'updated': True  # Маркер что данные обновлены
                        }
                        logging.info(f"Добавлена практика для группы {group_col}: {practice_info.strip()}")
                        continue
            except (TypeError, AttributeError) as e:
                logging.error(f"Ошибка при обработке практики для группы {group_col}: {e}")
                continue  # Пропускаем группу при ошибке
            #
            day_col = df.columns[0]
            cabinet_col = df.columns[df.columns.get_loc(group_col) + 1]
            current_day = None
            lesson_counter = 0
            week_lessons = {1: [], 2: []}
            i = 0
            while i < len(df):
                row = df.iloc[i]
                if i >= practice_start if len(practice_rows) > 0 else False:
                    break
                # Новый день недели
                if pd.notna(row[day_col]) and str(row[day_col]).strip():
                    if current_day and (week_lessons[1] or week_lessons[2]):
                        if current_day not in schedule_data[group_col]:
                            schedule_data[group_col][current_day] = {1: [], 2: []}
                        schedule_data[group_col][current_day][1].extend(week_lessons[1])
                        schedule_data[group_col][current_day][2].extend(week_lessons[2])
                    current_day = str(row[day_col]).strip()
                    week_lessons = {1: [], 2: []}
                    lesson_counter = 0
                time = str(row.get('Интервал', '')).strip()
                cell_value = str(row.get(group_col, '')).strip()
                cabinet_value = str(row.get(cabinet_col, '')).strip()
                # Пропуск пустых строк
                if not time or cell_value.lower() == 'nan':
                    i += 1
                    continue
                # Определяем номер пары
                if time and not any(x in time.lower() for x in ['снимаются', 'проводятся']):
                    lesson_counter += 1
                else:
                    i += 1
                    continue

                # Новый алгоритм: ищем разделитель '-----' и распределяем пары по неделям
                # Собираем блок пар для дня
                day_pairs = []
                day_cabinets = []
                day_times = []
                start_i = i
                while i < len(df):
                    row = df.iloc[i]
                    pair_value = str(row.get(group_col, '')).strip()
                    pair_cabinet = str(row.get(cabinet_col, '')).strip()
                    pair_time = str(row.get('Интервал', '')).strip()
                    if not pair_time or pair_value.lower() == 'nan':
                        i += 1
                        continue
                    if pair_value == "-----":
                        break
                    day_pairs.append(pair_value)
                    day_cabinets.append(pair_cabinet)
                    day_times.append(pair_time)
                    i += 1

                # Проверяем, есть ли разделитель '-----' в этом дне
                has_split = False
                split_index = None
                for idx in range(start_i, i):
                    row = df.iloc[idx]
                    if str(row.get(group_col, '')).strip() == "-----":
                        has_split = True
                        split_index = idx - start_i
                        break

                # Если есть разделитель, распределяем пары по неделям
                if has_split:
                    # Если '-----' над предметом (то есть split_index == 0)
                    if split_index == 0:
                        # 1 неделя — пары до разделителя, 2 неделя — после
                        for j in range(len(day_pairs)):
                            subject, teacher, room, subgroup = split_subject_teacher(day_pairs[j])
                            room_final = room if room else (day_cabinets[j] if day_cabinets[j] and day_cabinets[j].lower() != 'nan' else '—')
                            lesson_dict = {
                                'lesson_number': j+1,
                                'time': day_times[j],
                                'subject': subject,
                                'teacher': teacher,
                                'room': room_final,
                                'subgroup': subgroup,
                                'week_number': 2 if j >= split_index else 1,
                                'is_subgroup': bool(subgroup),
                                'file_hash': file_hash
                            }
                            if j < split_index:
                                week_lessons[1].append(lesson_dict)
                            else:
                                week_lessons[2].append(lesson_dict)
                    else:
                        # 1 неделя — пары до разделителя + предмет над '-----', 2 неделя — только пары до разделителя
                        for j in range(len(day_pairs)):
                            subject, teacher, room, subgroup = split_subject_teacher(day_pairs[j])
                            room_final = room if room else (day_cabinets[j] if day_cabinets[j] and day_cabinets[j].lower() != 'nan' else '—')
                            lesson_dict = {
                                'lesson_number': j+1,
                                'time': day_times[j],
                                'subject': subject,
                                'teacher': teacher,
                                'room': room_final,
                                'subgroup': subgroup,
                                'week_number': 1 if j <= split_index else 2,
                                'is_subgroup': bool(subgroup),
                                'file_hash': file_hash
                            }
                            if j <= split_index:
                                week_lessons[1].append(lesson_dict)
                            else:
                                week_lessons[2].append(lesson_dict)
                    i += 1  # пропускаем строку с '-----'
                else:
                    # Нет разделителя — обычная обработка
                    for j in range(len(day_pairs)):
                        subject, teacher, room, subgroup = split_subject_teacher(day_pairs[j])
                        room_final = room if room else (day_cabinets[j] if day_cabinets[j] and day_cabinets[j].lower() != 'nan' else '—')
                        lesson_dict = {
                            'lesson_number': j+1,
                            'time': day_times[j],
                            'subject': subject,
                            'teacher': teacher,
                            'room': room_final,
                            'subgroup': subgroup,
                            'week_number': 1,
                            'is_subgroup': bool(subgroup),
                            'file_hash': file_hash
                        }
                        week_lessons[1].append(lesson_dict)
            # Добавляем последний день
            if current_day and (week_lessons[1] or week_lessons[2]):
                if not isinstance(schedule_data[group_col], dict):
                    schedule_data[group_col] = {}
                if current_day not in schedule_data[group_col]:
                    schedule_data[group_col][current_day] = {1: [], 2: []}
                schedule_data[group_col][current_day][1].extend(week_lessons[1])
                schedule_data[group_col][current_day][2].extend(week_lessons[2])
        _schedule_cache = schedule_data
        _schedule_cache_hash = file_hash
        return schedule_data
    except Exception:
        return {}

def fetch_replacements():
    """Получает и парсит замены в расписании"""
    try:
        try:
            headers = get_random_headers()
            resp = requests.get(REPLACEMENTS_URL, headers=headers)
            resp.raise_for_status()
        except requests.exceptions.RequestException as e:
            print(f"Ошибка при получении файла замен: {e}")
            return {}

        if not resp.content:
            print("Получен пустой файл замен")
            return {}
            
        try:
            doc = Document(BytesIO(resp.content))
        except Exception as e:
            logging.error(f"Ошибка при открытии файла Word с заменами: {e}")
            return {}
        logging.info(f"Найдено таблиц в документе замен: {len(doc.tables)}")
        replacements_data = {}
        current_date = None
        if not doc.tables:
            print("В документе не найдено таблиц с заменами")
            return {}
            
        for table_idx, table in enumerate(doc.tables):
            try:
                logging.debug(f"Обработка таблицы замен {table_idx}, найдено строк: {len(table.rows)}")
                
                if not table.rows:
                    logging.warning(f"Таблица замен {table_idx} пуста")
                    continue
                    
                for row_idx, row in enumerate(table.rows):
                    try:
                        cells = [cell.text.strip() for cell in row.cells]
                        logging.debug(f"Обработка строки замен {row_idx}: {cells}")

                        # Поиск даты в первой или второй ячейке
                        date_candidate = None
                        for c in cells[:2]:
                            if c and "20" in c and "." in c and len(c) >= 8:
                                date_candidate = c
                                break
                        # Если нашли дату, обновляем current_date и пропускаем строку
                        if date_candidate:
                            current_date = date_candidate
                            print(f"Найдена дата: {current_date}")
                            continue

                        # Пропуск строк без даты и без группы
                        if not current_date:
                            # Если строка не содержит группу, это заголовок или пустая строка
                            if not cells or not any(cells):
                                continue
                            # Если первая ячейка не похожа на группу, пропускаем
                            group_candidate = cells[0].strip() if cells else ""
                            if not group_candidate or group_candidate.lower() in ["шифр группы", ""]:
                                continue
                            # Если нет даты, но есть группа, просто пропускаем (не спамим лог)
                            continue

                        # Пропуск пустых строк
                        if not any(cells):
                            continue

                        # Основная логика разбора замен
                        # Ожидается: [группа, № пары, № пары, дисциплина, ФИО, № пары, дисциплина, ФИО, аудитория]
                        # Но иногда бывает только 4 колонки: [группа, № пары, предмет, кабинет]
                        group = cells[0].strip() if len(cells) > 0 else ""
                        if not group:
                            continue

                        # Универсальный разбор: ищем все замены в строке (может быть 2 замены для одной группы)
                        # Пример: ['Бд-241', '3-4', '3-4', 'МДК.01.01', 'Литвинова', '3-4', 'История России', 'Лыкова', '401-1\n404-1']
                        # Первая замена: [1] пара, [3] предмет, [4] ФИО, [8] аудитория (если есть)
                        # Вторая замена: [5] пара, [6] предмет, [7] ФИО, [8] аудитория (если есть)
                        # Если только 4 колонки: [группа, пара, предмет, кабинет]

                        if group not in replacements_data:
                            replacements_data[group] = {}
                        if current_date not in replacements_data[group]:
                            replacements_data[group][current_date] = []

                        # Если строка длинная (две замены)
                        if len(cells) >= 8:
                            # Первая замена
                            lesson1 = cells[1].strip()
                            subject1 = cells[3].strip()
                            teacher1 = cells[4].strip()
                            room1 = cells[8].strip() if len(cells) > 8 else ""
                            if subject1 and lesson1:
                                replacements_data[group][current_date].append({
                                    'lesson': lesson1,
                                    'subject': subject1,
                                    'teacher': teacher1,
                                    'room': room1,
                                })
                                logging.debug(f"Добавлена первая замена для группы {group}")
                            # Вторая замена
                            lesson2 = cells[5].strip()
                            subject2 = cells[6].strip()
                            teacher2 = cells[7].strip()
                            room2 = cells[8].strip() if len(cells) > 8 else ""
                            if subject2 and lesson2:
                                replacements_data[group][current_date].append({
                                    'lesson': lesson2,
                                    'subject': subject2,
                                    'teacher': teacher2,
                                    'room': room2,
                                })
                                logging.debug(f"Добавлена вторая замена для группы {group}")
                        # Если строка обычная (одна замена)
                        elif len(cells) >= 4:
                            lesson = cells[1].strip()
                            subject = cells[2].strip()
                            teacher = cells[3].strip() if len(cells) > 3 else ""
                            room = cells[4].strip() if len(cells) > 4 else ""
                            if subject and lesson:
                                replacements_data[group][current_date].append({
                                    'lesson': lesson,
                                    'subject': subject,
                                    'teacher': teacher,
                                    'room': room,
                                })
                                print(f"Добавлена замена для группы {group}")
                        # Если строка короткая, пропускаем
                        else:
                            continue

                    except Exception as e:
                        print(f"Ошибка при обработке строки {row_idx}: {e}")
                        continue
                        
            except Exception as e:
                print(f"Ошибка при обработке таблицы {table_idx}: {e}")
                continue

        if not replacements_data:
            logging.warning("Не найдено данных о заменах")
        else:
            logging.info(f"Найдены замены для групп: {list(replacements_data.keys())}")
        return replacements_data
    except Exception as e:
        print(f"Ошибка при получении замен: {e}")
        return {}

def extract_groups_from_schedule():
    """Извлекает список групп из расписания"""
    try:
        schedule_data = fetch_schedule()
        if not schedule_data:
            print("Не удалось получить данные расписания")
            return []
            
        try:
            groups = list(schedule_data.keys())
            if not groups:
                return []
            cleaned_groups = []
            for group in groups:
                try:
                    group = str(group).strip()
                    if group and group not in ['Время', 'Дата', 'День', '']:
                        cleaned_groups.append(group)
                except (AttributeError, TypeError):
                    continue
            if not cleaned_groups:
                return []
            return sorted(list(set(cleaned_groups)))
        except Exception:
            return []
    except Exception:
        return []

# Для теста:
if __name__ == "__main__":
    schedule = fetch_schedule()
    for group, lessons in schedule.items():
        print(f"\nРасписание для группы {group}:")
        for day in lessons:
            print(format_day_schedule(lessons, day))
    replacements = fetch_replacements()
    print(replacements)
    groups = extract_groups_from_schedule()
    print("Найденные группы:", groups)
